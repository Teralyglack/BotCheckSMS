from __future__ import annotations

import textwrap
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from app.services.database import HistoryRecord


FORMAT_NAMES = {
    "transcript": "Дословная транскрипция",
    "summary": "Краткое содержание",
    "bullets": "Тезисы",
}


class ResultExporter:
    def __init__(self, export_dir: Path) -> None:
        self.export_dir = export_dir
        self.export_dir.mkdir(parents=True, exist_ok=True)

    def export_docx(self, record: HistoryRecord, messages: list[dict]) -> Path:
        path = self.export_dir / f"result_{record.request_id}.docx"
        doc = Document()
        doc.add_heading("Результат обработки переписки", level=1)
        doc.add_paragraph(f"Номер обработки: {record.request_id}")
        doc.add_paragraph(f"Формат: {FORMAT_NAMES.get(record.output_format, record.output_format)}")
        doc.add_paragraph(f"Количество сообщений: {record.message_count}")
        doc.add_paragraph(f"Дата: {record.created_at}")

        doc.add_heading("Итог", level=2)
        for paragraph in record.result_text.split("\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())

        if messages:
            doc.add_heading("Исходные сообщения", level=2)
            for index, item in enumerate(messages, start=1):
                doc.add_paragraph(
                    f"{index}. {item['role']} ({item['kind']}): {item['content']}",
                    style=None,
                )

        doc.save(path)
        return path

    def export_pdf(self, record: HistoryRecord, messages: list[dict]) -> Path:
        path = self.export_dir / f"result_{record.request_id}.pdf"
        font_name = self._register_font()
        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                name="AppNormal",
                parent=styles["Normal"],
                fontName=font_name,
                fontSize=10,
                leading=14,
            )
        )
        styles.add(
            ParagraphStyle(
                name="AppTitle",
                parent=styles["Title"],
                fontName=font_name,
                fontSize=17,
                leading=22,
                spaceAfter=14,
            )
        )
        styles.add(
            ParagraphStyle(
                name="AppHeading",
                parent=styles["Heading2"],
                fontName=font_name,
                fontSize=13,
                leading=16,
                spaceBefore=12,
                spaceAfter=8,
            )
        )

        story = [Paragraph("Результат обработки переписки", styles["AppTitle"])]
        meta = (
            f"Номер обработки: {record.request_id}<br/>"
            f"Формат: {FORMAT_NAMES.get(record.output_format, record.output_format)}<br/>"
            f"Количество сообщений: {record.message_count}<br/>"
            f"Дата: {record.created_at}"
        )
        story.append(Paragraph(meta, styles["AppNormal"]))
        story.append(Spacer(1, 10))
        story.append(Paragraph("Итог", styles["AppHeading"]))
        self._append_text(story, record.result_text, styles["AppNormal"])

        if messages:
            story.append(Paragraph("Исходные сообщения", styles["AppHeading"]))
            for index, item in enumerate(messages, start=1):
                text = f"{index}. {item['role']} ({item['kind']}): {item['content']}"
                self._append_text(story, text, styles["AppNormal"])

        doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        doc.build(story)
        return path

    @staticmethod
    def _append_text(story: list, text: str, style: ParagraphStyle) -> None:
        for raw_paragraph in text.split("\n"):
            paragraph = raw_paragraph.strip()
            if not paragraph:
                story.append(Spacer(1, 6))
                continue
            # ReportLab Paragraph cannot safely accept arbitrary user text without escaping.
            for part in textwrap.wrap(paragraph, width=170, replace_whitespace=False) or [paragraph]:
                story.append(Paragraph(escape(part), style))
            story.append(Spacer(1, 4))

    @staticmethod
    def _register_font() -> str:
        candidates = [
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/calibri.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf"),
        ]
        for candidate in candidates:
            if candidate.exists():
                pdfmetrics.registerFont(TTFont("AppFont", str(candidate)))
                return "AppFont"
        # PDF will still be generated, but Cyrillic may display incorrectly without a Unicode font.
        return "Helvetica"
